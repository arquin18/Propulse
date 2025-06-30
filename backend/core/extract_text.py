"""
Text Extraction Utilities
Handles parsing of PDF and DOCX documents with chunking capabilities.
"""

import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

import pypdf
from docx import Document
from sentence_transformers import SentenceTransformer
import numpy as np
from loguru import logger


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""
    content: str
    source_file: str
    chunk_id: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]


class TextExtractor:
    """Extract and chunk text from PDF and DOCX documents."""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize the text extractor.
        
        Args:
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of overlapping characters between chunks
            min_chunk_size: Minimum size for a chunk to be valid
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
    def extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                        
            logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    def extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_from_docx(file_path)
        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return self._clean_text(f.read())
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.
        
        Args:
            text: Raw text content
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\'\/\$\%\&\*\+\=\<\>\@\#]', '', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def chunk_text(
        self,
        text: str,
        source_file: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[TextChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            source_file: Source file name
            metadata: Additional metadata for chunks
            
        Returns:
            List of text chunks
        """
        if not text or len(text) < self.min_chunk_size:
            return []
        
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = min(start + self.chunk_size, len(text))
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_ends = []
                for i in range(max(end - 100, start), end):
                    if text[i] in '.!?':
                        sentence_ends.append(i + 1)
                
                if sentence_ends:
                    end = sentence_ends[-1]
            
            chunk_text = text[start:end].strip()
            
            if len(chunk_text) >= self.min_chunk_size:
                chunk = TextChunk(
                    content=chunk_text,
                    source_file=source_file,
                    chunk_id=chunk_id,
                    start_char=start,
                    end_char=end,
                    metadata=metadata or {}
                )
                chunks.append(chunk)
                chunk_id += 1
            
            # Move start position with overlap
            if end >= len(text):
                break
            start = max(start + 1, end - self.chunk_overlap)
        
        logger.info(f"Created {len(chunks)} chunks from {source_file}")
        return chunks
    
    def extract_and_chunk(
        self,
        file_path: Path,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[TextChunk]:
        """
        Extract text from file and split into chunks.
        
        Args:
            file_path: Path to the file
            metadata: Additional metadata for chunks
            
        Returns:
            List of text chunks
        """
        text = self.extract_text(file_path)
        source_file = file_path.name
        
        # Add file metadata
        file_metadata = {
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "file_extension": file_path.suffix,
            **(metadata or {})
        }
        
        return self.chunk_text(text, source_file, file_metadata) 